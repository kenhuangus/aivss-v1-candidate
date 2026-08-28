#!/usr/bin/env python3
"""AIVSS markdown→DOCX with leading prose before every figure and table."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

import doc_pipeline
from figure_table_intros import FIGURE_INTROS, TABLE_INTROS

BODY_PT = 11.0


def add_lead_prose(doc, text: str) -> None:
    """Insert a body paragraph; supports **bold** spans."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    for part in re.split(r"(\*\*.*?\*\*)", text):
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part:
            p.add_run(part)


def _figure_number(caption: str) -> str:
    m = re.match(r"Figure (\d+):", caption)
    return m.group(1) if m else "?"


def _insert_diagram(doc, sec_text: str, diagrams_dir: Path | None, diagram_map: dict, insert_figure) -> bool:
    for sec_key, (png_name, caption) in diagram_map.items():
        if sec_text.startswith(f"{sec_key} ") or sec_text.startswith(sec_key):
            if diagrams_dir:
                intro_tpl = FIGURE_INTROS.get(sec_key, "")
                if intro_tpl:
                    add_lead_prose(doc, intro_tpl.format(n=_figure_number(caption)))
                insert_figure(doc, diagrams_dir / png_name, caption)
            return True
    return False


def _insert_table(doc, table_md, assets_dir, counters, title, insert_table_fn):
    n = counters["table"] + 1
    cap_title = title or doc_pipeline._derive_table_title(table_md)
    intro_tpl = TABLE_INTROS.get(title) or TABLE_INTROS.get(cap_title, "")
    if intro_tpl:
        add_lead_prose(doc, intro_tpl.format(n=n, title=cap_title))
    insert_table_fn(doc, table_md, assets_dir, counters, title=title)


def markdown_to_docx(
    md_path: Path,
    docx_path: Path,
    diagrams_dir: Path | None = None,
    diagram_map: dict | None = None,
    insert_figure=None,
    insert_table_fn=None,
):
    insert_figure = insert_figure or doc_pipeline.insert_figure
    insert_table_fn = insert_table_fn or doc_pipeline.insert_table_as_image

    print(f"[*] Converting {md_path.name} to {docx_path.name} with 6.4-inch constrained layout...")
    content = md_path.read_text(encoding="utf-8")
    doc = Document()

    for s in doc.sections:
        s.page_width = Inches(8.5)
        s.page_height = Inches(11.0)
        s.top_margin = Inches(1.0)
        s.bottom_margin = Inches(1.0)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(BODY_PT)
    normal_style.font.color.rgb = RGBColor.from_string("1E293B")

    diagram_map = diagram_map or {}
    counters = {"table": 0}
    assets_tables_dir = md_path.parent / "assets" / "tables"
    pending_table_title = ""

    lines = content.splitlines()
    i = 0
    in_code_block = False
    lang = ""
    code_lines = []

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            if in_code_block:
                in_code_block = False
                doc_pipeline.add_themed_code_block(doc, "\n".join(code_lines), lang)
                code_lines = []
            else:
                in_code_block = True
                lang = line[3:].strip().lower()
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        if stripped.startswith("<!--"):
            m_t = re.match(r"^<!--\s*table:\s*(.*?)\s*-->$", stripped)
            if m_t:
                pending_table_title = m_t.group(1)
            while i < len(lines) and "-->" not in lines[i]:
                i += 1
            i += 1
            continue

        if line.startswith("> [!") and any(k in line for k in ["Diagram", "Pipeline", "Dynamics", "Workflow"]):
            while i < len(lines) and lines[i].startswith(">"):
                i += 1
            continue

        if line.startswith("# "):
            h = doc.add_heading(level=1)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
            run = h.add_run(line[2:].strip().replace("**", ""))
            run.font.name = "Arial"
            run.font.size = Pt(20)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string("0F172A")

        elif line.startswith("## "):
            sec_text = line[3:].strip().replace("**", "")
            h = doc.add_heading(level=2)
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run(sec_text)
            run.font.name = "Arial"
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string("1E3A8A")
            _insert_diagram(doc, sec_text, diagrams_dir, diagram_map, insert_figure)

        elif line.startswith("### "):
            sec_text = line[4:].strip().replace("**", "")
            h = doc.add_heading(level=3)
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(sec_text)
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor.from_string("334155")
            _insert_diagram(doc, sec_text, diagrams_dir, diagram_map, insert_figure)

        elif line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            i -= 1
            _insert_table(
                doc,
                "\n".join(table_lines),
                assets_tables_dir,
                counters,
                pending_table_title,
                insert_table_fn,
            )
            pending_table_title = ""

        elif line.startswith("> "):
            callout_lines = []
            while i < len(lines) and lines[i].startswith(">"):
                callout_lines.append(lines[i][2:].strip())
                i += 1
            i -= 1
            tbl = doc.add_table(rows=1, cols=1)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = tbl.cell(0, 0)
            doc_pipeline.set_cell_background(cell, "EFF6FF")
            doc_pipeline.set_cell_margins(cell, top=100, bottom=100, left=180, right=180)
            cell.width = Inches(6.4)
            cp = cell.paragraphs[0]
            cp.paragraph_format.space_before = Pt(2)
            cp.paragraph_format.space_after = Pt(2)
            run = cp.add_run(" ".join(callout_lines))
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor.from_string("1E3A8A")

        elif stripped and not stripped.startswith("---"):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            parts = re.split(r"(\*\*.*?\*\*)", stripped)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                elif part.startswith("`") and part.endswith("`"):
                    r = p.add_run(part[1:-1])
                    r.font.name = "Consolas"
                    r.font.size = Pt(9.5)
                else:
                    p.add_run(part)

        i += 1

    doc.save(docx_path)
    print(f"  [+] Saved DOCX: {docx_path.name} ({docx_path.stat().st_size // 1024} KB)")
    print(f"  [+] Data tables converted to images: {counters['table']} (native data tables: 0)")

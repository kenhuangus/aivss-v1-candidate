#!/usr/bin/env python3
"""Build pipeline with 12pt+ figure/table captions and large-font table infographics.

IMAGE STYLE RULE (project-wide, remember for all future docs):
  - All embedded images (figures + table infographics) use white (#ffffff) canvas backgrounds.
  - No dark or black fills anywhere inside images. Use light pastels + borders + dark text.
"""

from __future__ import annotations

import sys
from pathlib import Path

from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SCRIPTS = Path(__file__).resolve().parent
SKILL_SCRIPTS = Path(r"C:\Users\kenhu\.claude\skills\google-doc-gws\scripts")

sys.path.insert(0, str(SCRIPTS))
sys.path.insert(0, str(SKILL_SCRIPTS))

import doc_pipeline  # noqa: E402
import table_to_image_large as tti  # noqa: E402
from aivss_markdown_to_docx import markdown_to_docx as aivss_markdown_to_docx  # noqa: E402

# Route all table rendering through the large-font module.
doc_pipeline.render_table_to_png = tti.render_table_to_png

BODY_PT = 11.0
CAPTION_PT = 12.0


def insert_figure_large(doc, png_path: Path, caption: str, width_in: float = 6.4):
    if not Path(png_path).exists():
        print(f"  [!] Missing figure PNG, skipped: {png_path}")
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.space_before = Pt(14)
    p_img.paragraph_format.space_after = Pt(6)
    p_img.add_run().add_picture(str(png_path), width=Inches(width_in))

    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(16)
    run_cap = p_cap.add_run(caption)
    run_cap.italic = True
    run_cap.font.size = Pt(CAPTION_PT)
    run_cap.font.bold = True
    run_cap.font.color.rgb = RGBColor.from_string("1E293B")


doc_pipeline.insert_figure = insert_figure_large


def rasterize_svgs_large(diagrams_dir: Path, scale: float = 8.0):
    import cairosvg

    print(f"[*] Rasterizing SVGs at scale={scale} for large readable figures...")
    for svg_file in diagrams_dir.glob("*.svg"):
        png_file = diagrams_dir / f"{svg_file.stem}.png"
        cairosvg.svg2png(url=str(svg_file), write_to=str(png_file), scale=scale)
        print(f"  [+] {png_file.name} ({png_file.stat().st_size // 1024} KB)")


def _preprocess_markdown(md_path: Path) -> Path:
    """doc_pipeline supports ATX headings #–### only; #### would render as literal text."""
    text = md_path.read_text(encoding="utf-8")
    out: list[str] = []
    changed = False
    for line in text.splitlines():
        if line.startswith("#### "):
            changed = True
            title = line[5:].strip()
            if out and out[-1].strip():
                out.append("")
            out.append(f"**{title}**")
            out.append("")
        else:
            out.append(line)
    if not changed:
        return md_path
    tmp = md_path.with_name(f"{md_path.stem}.build{md_path.suffix}")
    tmp.write_text("\n".join(out) + "\n", encoding="utf-8")
    return tmp


def markdown_to_docx_large(md_path: Path, docx_path: Path, diagrams_dir: Path | None, diagram_map: dict):
    build_md = _preprocess_markdown(md_path)
    try:
        aivss_markdown_to_docx(
            build_md,
            docx_path,
            diagrams_dir,
            diagram_map,
            insert_figure=insert_figure_large,
            insert_table_fn=doc_pipeline.insert_table_as_image,
        )
    finally:
        if build_md != md_path and build_md.exists():
            build_md.unlink()
    from docx import Document

    doc = Document(docx_path)
    normal = doc.styles["Normal"]
    normal.font.size = Pt(BODY_PT)
    doc.save(docx_path)


def build(md_path: Path, docx_path: Path, diagrams_dir: Path, diagram_map: dict):
    rasterize_svgs_large(diagrams_dir)
    markdown_to_docx_large(md_path, docx_path, diagrams_dir, diagram_map)


def publish(file_id: str, docx_path: Path, qa_dir: Path) -> bool:
    return doc_pipeline.upload_and_qa_live_doc(file_id, docx_path, qa_dir)
